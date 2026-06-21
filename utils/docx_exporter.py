from datetime import date
from typing import Optional


class DOCXExporter:

    @staticmethod
    def _get_docx():
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            return Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH
        except ImportError:
            return None

    @staticmethod
    def is_available() -> bool:
        return DOCXExporter._get_docx() is not None

    @staticmethod
    def export_checkins(logs, path: str) -> bool:
        docx = DOCXExporter._get_docx()
        if not docx:
            return False
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = docx
        doc = Document()
        doc.add_heading("Rapport des check-ins", 0)
        doc.add_paragraph(f"Généré le {date.today()}").style = doc.styles["Normal"]
        if not logs:
            doc.add_paragraph("Aucune donnée de check-in.")
            doc.save(path)
            return True
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Date", "Sommeil", "Humeur", "Énergie", "Score", "Notes"]):
            hdr[i].text = label
        for log in logs:
            row = table.add_row().cells
            row[0].text = str(log.date)
            row[1].text = f"{getattr(log, 'sommeil', '-')}h"
            row[2].text = f"{getattr(log, 'humeur', '-')}/5"
            row[3].text = f"{getattr(log, 'energie', '-')}/5"
            row[4].text = str(getattr(log, 'score_productivite', '-'))
            row[5].text = (getattr(log, 'notes', "") or "")[:50]
        doc.save(path)
        return True

    @staticmethod
    def export_journal(entries, path: str) -> bool:
        docx = DOCXExporter._get_docx()
        if not docx:
            return False
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = docx
        doc = Document()
        doc.add_heading("Journal intime", 0)
        doc.add_paragraph(f"Généré le {date.today()}")
        if not entries:
            doc.add_paragraph("Aucune entrée de journal.")
            doc.save(path)
            return True
        for entry in entries:
            doc.add_heading(f"{entry.date}", level=2)
            p = doc.add_paragraph(getattr(entry, 'contenu', '(vide)'))
            doc.add_paragraph(f"Humeur: {getattr(entry, 'humeur', '-')}/5").italic = True
            doc.add_paragraph("")
        doc.save(path)
        return True

    @staticmethod
    def export_chat_history(messages: list, path: str) -> bool:
        docx = DOCXExporter._get_docx()
        if not docx:
            return False
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = docx
        doc = Document()
        doc.add_heading("Historique de conversation IA", 0)
        doc.add_paragraph(f"Généré le {date.today()}")
        if not messages:
            doc.add_paragraph("Aucun message.")
            doc.save(path)
            return True
        for msg in messages:
            role = msg.get("role", "inconnu").upper()
            content = msg.get("parts", [{}])[0].get("text", "") if isinstance(msg.get("parts"), list) else str(msg.get("parts", ""))
            doc.add_heading(role, level=2)
            doc.add_paragraph(content)
        doc.save(path)
        return True

    @staticmethod
    def export_sessions(sessions, path: str) -> bool:
        docx = DOCXExporter._get_docx()
        if not docx:
            return False
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = docx
        doc = Document()
        doc.add_heading("Sessions Pomodoro", 0)
        doc.add_paragraph(f"Généré le {date.today()}")
        if not sessions:
            doc.add_paragraph("Aucune session.")
            doc.save(path)
            return True
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Date", "Durée", "Sujet", "Statut"]):
            hdr[i].text = label
        for s in sessions:
            row = table.add_row().cells
            row[0].text = str(getattr(s, 'date', ''))
            row[1].text = f"{getattr(s, 'duree', 0)}min"
            row[2].text = getattr(s, 'sujet', '') or ""
            row[3].text = getattr(s, 'statut', '')
        doc.save(path)
        return True

    @staticmethod
    def export_goals(goals, path: str) -> bool:
        docx = DOCXExporter._get_docx()
        if not docx:
            return False
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = docx
        doc = Document()
        doc.add_heading("Objectifs", 0)
        doc.add_paragraph(f"Généré le {date.today()}")
        if not goals:
            doc.add_paragraph("Aucun objectif.")
            doc.save(path)
            return True
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Titre", "Catégorie", "Progrès", "Échéance"]):
            hdr[i].text = label
        for g in goals:
            row = table.add_row().cells
            row[0].text = getattr(g, 'titre', '')
            row[1].text = getattr(g, 'categorie', '')
            row[2].text = f"{getattr(g, 'progres', 0)}/{getattr(g, 'cible', 100)}"
            row[3].text = str(getattr(g, 'date_limite', ''))
        doc.save(path)
        return True
